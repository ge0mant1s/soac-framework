"""
Operational Model Parser - Parses security operational model DOCX files
and extracts structured attack patterns, phases, detection rules, and playbooks
"""
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
import docx
from docx import Document
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class OperationalModelParser:
    """
    Parses operational model documents and extracts:
    - Attack phases and patterns
    - Detection queries
    - Correlation rules
    - MITRE ATT&CK mappings
    - Response playbooks
    - KPI metrics
    """
    
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.model_data = {
            "id": "",
            "name": "",
            "version": "1.0",
            "created_at": datetime.utcnow().isoformat(),
            "objective": {},
            "correlation_pattern": {},
            "detection_queries": {},
            "alert_policy": {},
            "dashboard_panels": [],
            "playbooks": [],
            "decision_matrix": [],
            "kpi_metrics": [],
            "strategic_impact": [],
            "deliverables": []
        }
    
    def parse(self) -> Dict[str, Any]:
        """Parse the entire operational model document"""
        try:
            logger.info(f"Parsing operational model: {self.docx_path}")
            
            # Extract model name from filename
            self.model_data["name"] = self._extract_name_from_filename()
            self.model_data["id"] = self._generate_model_id()
            
            # Parse all sections
            full_text = self._get_full_text()
            
            self.model_data["objective"] = self._parse_objective(full_text)
            self.model_data["correlation_pattern"] = self._parse_correlation_pattern(full_text)
            self.model_data["detection_queries"] = self._parse_detection_queries(full_text)
            self.model_data["alert_policy"] = self._parse_alert_policy(full_text)
            self.model_data["dashboard_panels"] = self._parse_dashboard_panels(full_text)
            self.model_data["playbooks"] = self._parse_playbooks(full_text)
            self.model_data["decision_matrix"] = self._parse_decision_matrix(full_text)
            self.model_data["kpi_metrics"] = self._parse_kpi_metrics(full_text)
            
            logger.info(f"Successfully parsed model: {self.model_data['name']}")
            return self.model_data
            
        except Exception as e:
            logger.error(f"Error parsing operational model: {str(e)}")
            raise
    
    def _get_full_text(self) -> str:
        """Extract all text from document"""
        paragraphs = [para.text for para in self.doc.paragraphs]
        
        # Also extract text from tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        
        return "\n".join(paragraphs)
    
    def _extract_name_from_filename(self) -> str:
        """Extract model name from filename"""
        name = self.docx_path.stem
        # Remove 'operational model' and 'Operational Model' variations
        name = re.sub(r'\s*operational\s+model\s*', '', name, flags=re.IGNORECASE)
        return name.strip().title()
    
    def _generate_model_id(self) -> str:
        """Generate unique model ID"""
        name = self.model_data["name"].upper().replace(" ", "_")
        return f"{name}_MODEL"
    
    def _parse_objective(self, text: str) -> Dict[str, str]:
        """Parse objective section"""
        objective = {"goal": "", "business_outcome": ""}
        
        # Find Goal
        goal_match = re.search(r'Goal:\s*(.+?)(?=Business Outcome:|$)', text, re.DOTALL | re.IGNORECASE)
        if goal_match:
            objective["goal"] = goal_match.group(1).strip()
        
        # Find Business Outcome
        outcome_match = re.search(r'Business Outcome:\s*(.+?)(?=\d+\.\s+Correlation Pattern|$)', text, re.DOTALL | re.IGNORECASE)
        if outcome_match:
            objective["business_outcome"] = outcome_match.group(1).strip()
        
        return objective
    
    def _parse_correlation_pattern(self, text: str) -> Dict[str, Any]:
        """Parse correlation pattern section"""
        pattern = {
            "pattern_id": "",
            "description": "",
            "phases": [],
            "correlation_window": "",
            "pivot_entities": []
        }
        
        # Extract pattern ID and description
        pattern_match = re.search(r'Pattern\s+(\w+)\s+[-—]\s+(.+?)(?=Phase|$)', text, re.DOTALL)
        if pattern_match:
            pattern["pattern_id"] = pattern_match.group(1).strip()
            pattern["description"] = pattern_match.group(2).strip()
        
        # Extract correlation window
        window_match = re.search(r'Correlation Window:\s*(.+?)(?=Pivot|$)', text, re.IGNORECASE)
        if window_match:
            pattern["correlation_window"] = window_match.group(1).strip()
        
        # Extract pivot entities
        pivot_match = re.search(r'Pivot Entities:\s*(.+?)(?=\d+\.\s+|$)', text, re.IGNORECASE)
        if pivot_match:
            entities_text = pivot_match.group(1).strip()
            pattern["pivot_entities"] = [e.strip() for e in entities_text.split(',')]
        
        # Parse phases from table (if present)
        pattern["phases"] = self._extract_phases_from_tables()
        
        return pattern
    
    def _extract_phases_from_tables(self) -> List[Dict[str, Any]]:
        """Extract attack phases from correlation pattern table"""
        phases = []
        
        for table in self.doc.tables:
            # Look for tables with headers: Phase, Source, Event Type/Indicator, Fields
            if len(table.rows) > 1:
                header_row = table.rows[0]
                headers = [cell.text.lower().strip() for cell in header_row.cells]
                
                if 'phase' in headers or 'source' in headers:
                    for row in table.rows[1:]:
                        if len(row.cells) >= 3:
                            phase = {
                                "name": row.cells[0].text.strip(),
                                "source": row.cells[1].text.strip() if len(row.cells) > 1 else "",
                                "indicators": row.cells[2].text.strip() if len(row.cells) > 2 else "",
                                "correlation_fields": row.cells[3].text.strip() if len(row.cells) > 3 else ""
                            }
                            if phase["name"]:  # Only add if phase name exists
                                phases.append(phase)
        
        return phases
    
    def _parse_detection_queries(self, text: str) -> Dict[str, Any]:
        """Parse detection queries section"""
        queries = {}
        
        # Find all Query A, B, C, D, E patterns
        query_pattern = r'Query\s+([A-E])\s+[-—]\s+([^\n]+)\s*\n(#[^\n]+(?:\n[^\n]+)*?)(?=Query\s+[A-E]|Combined|$)'
        matches = re.finditer(query_pattern, text, re.DOTALL)
        
        for match in matches:
            query_id = match.group(1)
            query_name = match.group(2).strip()
            query_text = match.group(3).strip()
            
            queries[f"query_{query_id}"] = {
                "id": query_id,
                "name": query_name,
                "query": query_text,
                "enabled": True
            }
        
        # Extract combined correlation rule
        combined_match = re.search(r'Combined Correlation Rule[^\n]*\n(join\([^\)]+\).*?)(?=Trigger:|Alert trigger:|$)', text, re.DOTALL | re.IGNORECASE)
        if combined_match:
            queries["combined_rule"] = {
                "query": combined_match.group(1).strip(),
                "type": "correlation"
            }
        
        # Extract trigger condition
        trigger_match = re.search(r'(?:Trigger:|Alert trigger:)\s*(.+?)(?=\d+\.\s+Alert Policy|$)', text, re.DOTALL | re.IGNORECASE)
        if trigger_match:
            queries["trigger_condition"] = trigger_match.group(1).strip()
        
        return queries
    
    def _parse_alert_policy(self, text: str) -> Dict[str, Any]:
        """Parse alert policy section"""
        policy = {
            "severity": "",
            "trigger_condition": "",
            "suppression_window": "",
            "escalation_path": "",
            "runbook_reference": ""
        }
        
        # Extract from Alert Policy table or section
        severity_match = re.search(r'Severity\s+(\w+)', text, re.IGNORECASE)
        if severity_match:
            policy["severity"] = severity_match.group(1).strip()
        
        trigger_match = re.search(r'Trigger(?:\s+Condition)?\s+(.+?)(?=Suppression|$)', text, re.IGNORECASE | re.DOTALL)
        if trigger_match:
            policy["trigger_condition"] = trigger_match.group(1).strip()
        
        suppression_match = re.search(r'Suppression(?:\s+Window)?\s+(.+?)(?=Escalation|$)', text, re.IGNORECASE)
        if suppression_match:
            policy["suppression_window"] = suppression_match.group(1).strip()
        
        escalation_match = re.search(r'Escalation Path\s+(.+?)(?=Runbook|$)', text, re.IGNORECASE)
        if escalation_match:
            policy["escalation_path"] = escalation_match.group(1).strip()
        
        runbook_match = re.search(r'Runbook\s+Reference\s+(.+?)(?=\d+\.\s+|$)', text, re.IGNORECASE)
        if runbook_match:
            policy["runbook_reference"] = runbook_match.group(1).strip()
        
        return policy
    
    def _parse_dashboard_panels(self, text: str) -> List[Dict[str, str]]:
        """Parse dashboard panels section"""
        panels = []
        
        # Extract from Dashboard Panels table
        for table in self.doc.tables:
            if len(table.rows) > 1:
                header_row = table.rows[0]
                headers = [cell.text.lower().strip() for cell in header_row.cells]
                
                if 'panel' in ' '.join(headers) or 'source' in headers:
                    for row in table.rows[1:]:
                        if len(row.cells) >= 3:
                            panel = {
                                "name": row.cells[0].text.strip(),
                                "source": row.cells[1].text.strip(),
                                "purpose": row.cells[2].text.strip()
                            }
                            if panel["name"]:
                                panels.append(panel)
        
        return panels
    
    def _parse_playbooks(self, text: str) -> List[Dict[str, Any]]:
        """Parse SOAR playbooks section"""
        playbooks = []
        
        # Find all PLAYBOOK sections
        playbook_pattern = r'PLAYBOOK\s+(\d+|[A-Z]+)[:\s—-]+([^\n]+)\n(.*?)(?=PLAYBOOK\s+\d+|PLAYBOOK\s+[A-Z]+|\d+\.\s+DECISION MATRIX|$)'
        matches = re.finditer(playbook_pattern, text, re.DOTALL | re.IGNORECASE)
        
        for match in matches:
            playbook_id = match.group(1).strip()
            playbook_name = match.group(2).strip()
            playbook_content = match.group(3).strip()
            
            steps = self._extract_playbook_steps(playbook_content)
            
            playbooks.append({
                "id": f"PB_{playbook_id}",
                "name": playbook_name,
                "steps": steps,
                "enabled": True
            })
        
        return playbooks
    
    def _extract_playbook_steps(self, content: str) -> List[Dict[str, str]]:
        """Extract steps from playbook content"""
        steps = []
        
        # Look for Step patterns (1, 2, 3... or emoji patterns)
        step_pattern = r'(?:Step\s+)?(?:\d+|[①②③④⃣])\s+(.+?)(?=(?:\d+|[①②③④⃣])|$)'
        matches = re.finditer(step_pattern, content, re.DOTALL)
        
        step_num = 1
        for match in matches:
            step_text = match.group(1).strip()
            if step_text:
                # Extract action, integration, logic
                parts = step_text.split('\n')
                action = parts[0] if len(parts) > 0 else step_text
                
                steps.append({
                    "step": step_num,
                    "action": action,
                    "description": step_text
                })
                step_num += 1
        
        return steps
    
    def _parse_decision_matrix(self, text: str) -> List[Dict[str, Any]]:
        """Parse decision matrix section"""
        matrix = []
        
        # Extract from Decision Matrix table
        for table in self.doc.tables:
            if len(table.rows) > 1:
                header_row = table.rows[0]
                headers = [cell.text.lower().strip() for cell in header_row.cells]
                
                if 'condition' in headers or 'response' in headers or 'action' in headers:
                    for row in table.rows[1:]:
                        if len(row.cells) >= 2:
                            decision = {
                                "condition": row.cells[0].text.strip(),
                                "response_path": row.cells[1].text.strip(),
                                "playbooks_triggered": row.cells[2].text.strip() if len(row.cells) > 2 else ""
                            }
                            if decision["condition"]:
                                matrix.append(decision)
        
        return matrix
    
    def _parse_kpi_metrics(self, text: str) -> List[Dict[str, str]]:
        """Parse KPI metrics section"""
        metrics = []
        
        # Extract from KPI table
        for table in self.doc.tables:
            if len(table.rows) > 1:
                header_row = table.rows[0]
                headers = [cell.text.lower().strip() for cell in header_row.cells]
                
                if 'metric' in headers or 'target' in headers:
                    for row in table.rows[1:]:
                        if len(row.cells) >= 2:
                            metric = {
                                "name": row.cells[0].text.strip(),
                                "target": row.cells[1].text.strip(),
                                "description": row.cells[2].text.strip() if len(row.cells) > 2 else ""
                            }
                            if metric["name"]:
                                metrics.append(metric)
        
        return metrics
    
    def save_to_json(self, output_path: Optional[str] = None) -> str:
        """Save parsed model to JSON file"""
        if not output_path:
            output_path = str(self.docx_path.parent / f"{self.model_data['id']}.json")
        
        with open(output_path, 'w') as f:
            json.dump(self.model_data, f, indent=2)
        
        logger.info(f"Saved operational model to: {output_path}")
        return output_path


def parse_operational_model(docx_path: str, output_dir: Optional[str] = None) -> Dict[str, Any]:
    """
    Convenience function to parse an operational model DOCX file
    
    Args:
        docx_path: Path to the DOCX file
        output_dir: Optional output directory for JSON file
    
    Returns:
        Parsed model data dictionary
    """
    parser = OperationalModelParser(docx_path)
    model_data = parser.parse()
    
    if output_dir:
        output_path = Path(output_dir) / f"{model_data['id']}.json"
        parser.save_to_json(str(output_path))
    
    return model_data
